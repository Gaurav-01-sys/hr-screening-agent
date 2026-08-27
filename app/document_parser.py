from __future__ import annotations

import base64
import asyncio
import json
import os
from argparse import Namespace
from functools import lru_cache
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import List, Optional
from urllib import error as urllib_error
from urllib.parse import urlparse
from urllib import request as urllib_request

from .config import load_env_file


OCRFLUX_PAGE_PROMPT = (
    "Below is the image of one page of a document. "
    "Just return the plain text representation of this document as if you were reading it naturally.\n"
    "ALL tables should be presented in HTML format.\n"
    "If there are images or figures in the page, present them as "
    '\"<Image>(left,top),(right,bottom)</Image>\", '
    "(left,top,right,bottom) are the coordinates of the top-left and "
    "bottom-right corners of the image or figure.\n"
    "Present all titles and headings as H1 headings.\n"
    "Do not hallucinate.\n"
)


def _normalize_text(text: str) -> str:
    cleaned_lines = [line.rstrip() for line in text.splitlines()]
    normalized: List[str] = []
    previous_blank = False

    for line in cleaned_lines:
        is_blank = not line.strip()
        if is_blank and previous_blank:
            continue
        normalized.append(line)
        previous_blank = is_blank

    return "\n".join(normalized).strip()


def _configured_value(name: str, default: str = "") -> str:
    # Load .env lazily so document parsing works in both FastAPI and Streamlit.
    load_env_file()
    return os.getenv(name, default).strip()


def _configured_int(name: str, default: int) -> int:
    value = _configured_value(name)
    try:
        return max(1, int(value)) if value else default
    except ValueError:
        return default


def _ocrflux_completion_url() -> str:
    configured_url = _configured_value("OCRFLUX_URL").rstrip("/")
    if not configured_url:
        return ""
    if configured_url.endswith("/v1/chat/completions"):
        return configured_url
    return f"{configured_url}/v1/chat/completions"


def _content_to_text(content: object) -> str:
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        chunks = []
        for item in content:
            if isinstance(item, dict) and isinstance(item.get("text"), str):
                chunks.append(item["text"])
        return "\n".join(chunks).strip()
    return ""


def _extract_natural_text(content: object) -> str:
    text = _content_to_text(content)
    if not text:
        return ""

    candidate = text.strip()
    if candidate.startswith("```") and candidate.endswith("```"):
        candidate = candidate.split("\n", 1)[-1]
        candidate = candidate.rsplit("```", 1)[0].strip()

    try:
        payload = json.loads(candidate)
    except json.JSONDecodeError:
        return _normalize_text(text)

    if isinstance(payload, dict) and isinstance(payload.get("natural_text"), str):
        return _normalize_text(payload["natural_text"])
    return _normalize_text(text)


def _render_pdf_pages(file_bytes: bytes) -> List[bytes]:
    """Render PDF pages for OCRFlux's image-based completion endpoint."""
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise RuntimeError(
            "OCRFlux server mode requires pypdfium2; install the project requirements."
        ) from exc

    document = pdfium.PdfDocument(file_bytes)
    rendered_pages: List[bytes] = []
    try:
        for page_index in range(len(document)):
            page = document[page_index]
            bitmap = page.render(scale=2.0)
            image = bitmap.to_pil().convert("RGB")
            buffer = BytesIO()
            image.save(buffer, format="PNG")
            rendered_pages.append(buffer.getvalue())
            bitmap.close()
            page.close()
    finally:
        document.close()
    return rendered_pages


def _ocrflux_remote_page(image_bytes: bytes) -> str:
    completion_url = _ocrflux_completion_url()
    if not completion_url:
        return ""

    encoded_image = base64.b64encode(image_bytes).decode("ascii")
    payload = {
        "model": _configured_value("OCRFLUX_MODEL", "ChatDOC/OCRFlux-3B"),
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": OCRFLUX_PAGE_PROMPT},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{encoded_image}"},
                    },
                ],
            }
        ],
        "temperature": 0.0,
    }
    headers = {"Content-Type": "application/json"}
    api_key = _configured_value("OCRFLUX_API_KEY")
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    request = urllib_request.Request(
        completion_url,
        data=json.dumps(payload).encode("utf-8"),
        headers=headers,
        method="POST",
    )
    timeout = _configured_int("OCRFLUX_TIMEOUT_SECONDS", 180)
    try:
        with urllib_request.urlopen(request, timeout=timeout) as response:
            response_payload = json.loads(response.read().decode("utf-8"))
    except (OSError, urllib_error.URLError, json.JSONDecodeError) as exc:
        print(f"OCRFlux server extraction failed: {exc}")
        return ""

    if not isinstance(response_payload, dict):
        return ""
    choices = response_payload.get("choices")
    if not isinstance(choices, list) or not choices:
        return ""
    first_choice = choices[0]
    if not isinstance(first_choice, dict):
        return ""
    message = first_choice.get("message")
    if not isinstance(message, dict):
        return ""
    return _extract_natural_text(message.get("content"))


def _extract_with_ocrflux_server(file_bytes: bytes) -> str:
    official_result = _extract_with_ocrflux_official_client(file_bytes)
    if official_result:
        return official_result

    try:
        pages = _render_pdf_pages(file_bytes)
    except Exception as exc:
        print(f"OCRFlux page rendering failed: {exc}")
        return ""

    page_texts = [_ocrflux_remote_page(page) for page in pages]
    usable_text = [text for text in page_texts if text]
    return _normalize_text("\n\n".join(usable_text))


def _extract_with_ocrflux_official_client(file_bytes: bytes) -> str:
    """Use OCRFlux's complete online pipeline when its optional package is installed."""
    try:
        from ocrflux.client import request as ocrflux_request
    except ImportError:
        return ""

    parsed_url = urlparse(_ocrflux_completion_url())
    if parsed_url.scheme != "http" or not parsed_url.hostname:
        # OCRFlux's bundled client opens a plain TCP connection. HTTPS uses the
        # lightweight urllib path below instead.
        return ""

    temporary_path: Optional[str] = None
    try:
        with NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
            handle.write(file_bytes)
            temporary_path = handle.name
        args = Namespace(
            model=_configured_value("OCRFLUX_MODEL", "ChatDOC/OCRFlux-3B"),
            skip_cross_page_merge=False,
            max_page_retries=_configured_int("OCRFLUX_MAX_PAGE_RETRIES", 1),
            url=f"{parsed_url.scheme}://{parsed_url.hostname}",
            port=parsed_url.port or 80,
        )
        result = asyncio.run(ocrflux_request(args, temporary_path))
    except Exception as exc:
        print(f"OCRFlux online pipeline failed: {exc}")
        return ""
    finally:
        if temporary_path:
            try:
                Path(temporary_path).unlink(missing_ok=True)
            except OSError:
                pass

    if isinstance(result, dict) and isinstance(result.get("document_text"), str):
        return _normalize_text(result["document_text"])
    return ""


@lru_cache(maxsize=1)
def _load_local_ocrflux_model(model_path: str) -> object:
    from vllm import LLM

    gpu_memory_utilization = float(
        _configured_value("OCRFLUX_GPU_MEMORY_UTILIZATION", "0.8")
    )
    return LLM(
        model=model_path,
        gpu_memory_utilization=gpu_memory_utilization,
        max_model_len=_configured_int("OCRFLUX_MAX_MODEL_CONTEXT", 8192),
    )


def _extract_with_ocrflux_local(file_name: str, file_bytes: bytes) -> str:
    model_path = _configured_value("OCRFLUX_MODEL_PATH")
    if not model_path:
        return ""

    try:
        from ocrflux.inference import parse
    except ImportError:
        print(
            "OCRFlux local mode is configured but not installed. "
            "Install requirements-ocrflux.txt in a GPU environment."
        )
        return ""

    temporary_path: Optional[str] = None
    try:
        with NamedTemporaryFile(
            suffix=Path(file_name).suffix or ".pdf", delete=False
        ) as handle:
            handle.write(file_bytes)
            temporary_path = handle.name
        result = parse(
            _load_local_ocrflux_model(model_path),
            temporary_path,
            max_page_retries=_configured_int("OCRFLUX_MAX_PAGE_RETRIES", 1),
        )
    except Exception as exc:
        print(f"OCRFlux local extraction failed: {exc}")
        return ""
    finally:
        if temporary_path:
            try:
                Path(temporary_path).unlink(missing_ok=True)
            except OSError:
                pass

    if isinstance(result, dict) and isinstance(result.get("document_text"), str):
        return _normalize_text(result["document_text"])
    return ""


def _extract_with_ocrflux(file_name: str, file_bytes: bytes) -> str:
    """Use OCRFlux when configured, preferring its hosted API over local GPU mode."""
    if _ocrflux_completion_url():
        return _extract_with_ocrflux_server(file_bytes)
    return _extract_with_ocrflux_local(file_name, file_bytes)


def _extract_pdf_text(file_name: str, file_bytes: bytes) -> str:
    # OCRFlux is the primary parser when a server or local GPU model is configured.
    ocrflux_text = _extract_with_ocrflux(file_name, file_bytes)
    if ocrflux_text:
        return ocrflux_text

    # Keep a lightweight fallback for normal text-based PDFs and CPU deployments.
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        parts = [page.extract_text() or "" for page in reader.pages]
        return _normalize_text("\n\n".join(parts))
    except Exception as exc:
        print(f"pypdf extraction failed: {exc}")
        return ""


def _extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(file_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    return _normalize_text("\n".join(parts))


def extract_uploaded_text(file_name: str, file_bytes: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(file_name, file_bytes)
    if suffix == ".docx":
        return _extract_docx_text(file_bytes)
    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")
